# -*- coding: utf-8 -*-
"""Build the DeepResearch resume project-entry DOCX (compact_reference_guide preset)."""

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap, qn
from docx.shared import Inches, Pt, RGBColor


OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "项目经历-DeepResearch多Agent深度研究智能助手.docx")
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
BORDER = "9AA5B1"
HEADER_FILL = "E8EEF5"
LATIN_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"


def set_run_font(run, size=11, bold=False, color=None, latin=LATIN_FONT, cjk=CJK_FONT):
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)


def style_font(style, size=11, bold=False, color=None, latin=LATIN_FONT, cjk=CJK_FONT):
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)


def set_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.makeelement(qn("w:numPr"), {})
    ilvl = ppr.makeelement(qn("w:ilvl"), {qn("w:val"): "0"})
    numid = ppr.makeelement(qn("w:numId"), {qn("w:val"): str(num_id)})
    numpr.append(ilvl)
    numpr.append(numid)
    ppr.append(numpr)


def wxml(fragment):
    """Parse a WordprocessingML fragment with the w namespace declared."""
    return parse_xml(f'<w:frag xmlns:w="{W_NS}">{fragment}</w:frag>')


def set_table_geometry(table, widths_dxa, indent_dxa=120, cell_margins=(80, 80, 120, 120)):
    tbl = table._tbl
    tblpr = tbl.tblPr
    # Width + indent
    tblpr.append(wxml(f'<w:tblW w:type="dxa" w:w="{sum(widths_dxa)}"/>'))
    tblpr.append(wxml(f'<w:tblInd w:type="dxa" w:w="{indent_dxa}"/>'))
    top, bottom, start, end = cell_margins
    tblpr.append(wxml(
        f'<w:tblCellMar w:top="{top}" w:left="{start}" '
        f'w:bottom="{bottom}" w:right="{end}"/>'
    ))
    # Grid
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = wxml(
        "<w:tblGrid>"
        + "".join(f'<w:gridCol w:w="{w}"/>' for w in widths_dxa)
        + "</w:tblGrid>"
    )
    tbl.insert(list(tbl).index(tblpr) + 1, grid)
    # Per-cell widths + vertical centering
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcpr.append(wxml(f'<w:tcW w:type="dxa" w:w="{widths_dxa[col_idx]}"/>'))
            valign = tcpr.makeelement(qn("w:vAlign"), {qn("w:val"): "center"})
            tcpr.append(valign)


def set_table_borders(table):
    tblpr = table._tbl.tblPr
    borders = wxml(
        "<w:tblBorders>"
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{BORDER}"/>'
        "</w:tblBorders>"
    )
    tblpr.append(borders)


def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    tcpr.append(wxml(f'<w:shd w:val="clear" w:color="auto" w:fill="{fill}"/>'))


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION.NEW_PAGE

    # Styles ----------------------------------------------------------------
    normal = doc.styles["Normal"]
    style_font(normal, 11, False, None)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    h1 = doc.styles["Heading 1"]
    style_font(h1, 16, True, BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    style_font(h2, 13, True, BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    style_font(bullet, 11, False, None)
    bullet.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    bullet.paragraph_format.line_spacing = 1.25
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)

    # Numbering (real bullet definition) ------------------------------------
    numbering = doc.part.numbering_part
    numbering_el = numbering.element
    for tag in ("w:abstractNum", "w:num"):
        for old in numbering_el.findall(qn(tag)):
            numbering_el.remove(old)
    numbering_el.append(wxml(
        '<w:abstractNum w:abstractNumId="0">'
        '<w:multiLevelType w:val="hybridMultilevel"/>'
        '<w:lvl w:ilvl="0">'
        '<w:start w:val="1"/>'
        '<w:numFmt w:val="bullet"/>'
        '<w:lvlText w:val="\u2022"/>'
        '<w:lvlJc w:val="left"/>'
        '<w:pPr><w:ind w:left="540" w:hanging="270"/></w:pPr>'
        '<w:rPr>'
        f'<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:eastAsia="{CJK_FONT}"/>'
        '<w:sz w:val="22"/>'
        "</w:rPr>"
        "</w:lvl>"
        "</w:abstractNum>"
    ))
    numbering_el.append(wxml(
        '<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
    ))

    # Title -----------------------------------------------------------------
    title = doc.add_paragraph("DeepResearch 企业级多 Agent 深度研究智能助手", style="Heading 1")
    for run in title.runs:
        set_run_font(run, 16, True, BLUE)

    tech = doc.add_paragraph()
    label = tech.add_run("技术栈：")
    set_run_font(label, 11, True, DARK_BLUE)
    value = tech.add_run(
        "Python、FastAPI、LangGraph、LangChain、Milvus、PostgreSQL、Redis、"
        "DashScope（Qwen 推理/Embedding）、Bocha Web Search、Vue3 + TypeScript + Vite、Docker Compose"
    )
    set_run_font(value, 11, False, None)

    # 项目介绍 ---------------------------------------------------------------
    doc.add_paragraph("项目介绍", style="Heading 2")
    intro = doc.add_paragraph(
        "面向企业级深度研究场景的多 Agent 智能研究平台，解决分析师在行业调研、竞品分析、政策解读等场景中"
        "跨多平台收集资料耗时且易遗漏、大模型直接回答存在幻觉与引用不可追溯等痛点。基于 LangGraph 搭建完整的"
        "多 Agent 协作链路：意图路由自动识别研究型查询，Planner 将大问题拆解为子问题，Web Scout 与 Local Scout "
        "并行检索网络与本地知识库，Evidence Judge 完成证据评分、去重与冲突检测，Analyst 评估证据完备性，Reflect "
        "针对信息缺口自动补搜，Writer 基于证据池撰写 2000-3000 字带引用角标的深度研报，并配套三层记忆体系与 "
        "Web/CLI 双入口，实现从“人找信息”到“AI 自动研究”的闭环。"
    )
    for run in intro.runs:
        set_run_font(run, 11, False, None)

    # 项目职责 ---------------------------------------------------------------
    doc.add_paragraph("项目职责", style="Heading 2")
    duties = [
        "负责多 Agent 编排与意图分流：基于 LangGraph 状态机构建 8 大专家 Agent（意图路由/规划/网络侦察/本地侦察/"
        "证据裁判/分析/反思/撰稿），用 TypedDict + Annotated 定义共享 ResearchState，实现节点间类型安全的数据流转；"
        "意图路由采用规则引擎 + LLM 双模态，将闲聊、简单问答与深度研究分流，减少约 60% 无效深度调用；通过条件边与"
        "迭代计数实现“分析→补搜→再检索”闭环，并用最大迭代次数防止死循环。",
        "负责双源并行检索与证据审计：基于 Bocha Web Search + Milvus RAG 构建双路检索，利用 LangGraph 异步并发并行"
        "执行网络与本地侦察，检索阶段耗时降低约 35%；证据裁判按信源类型评分（本地 0.92、官方 0.88、主流媒体 0.72、"
        "普通站点 0.58），完成去重、冲突检测与审计标记，低质量信源占比由 45% 降至 12%；对 snippet 截断与结构化剪枝，"
        "控制大模型上下文 Token 规模。",
        "负责反思补搜与引用溯源：Analyst 节点强制评估证据完备性并输出信息缺口，Reflect 节点针对缺口生成补充查询触发"
        "二次检索，研究完备性满意度由 62% 提升至 89%；Writer 仅允许使用 source_index 中的合法来源编号，通过正则校验"
        "自动移除幻觉引用并渲染网络/本地分类参考列表，在 200 条内部测评集上幻觉率由约 25% 降至 6%、引用准确率达 94%。",
        "负责记忆系统与服务化：设计短期（会话）、长期（用户画像/任务）、语义（Milvus 向量）三层记忆，短期消息超阈值"
        "触发滚动摘要压缩并物理清理旧消息；Milvus 异常时自动降级 PostgreSQL ILIKE 检索，按 tenant_id/user_id 强隔离并"
        "丢弃跨用户命中，避免记忆串户；基于 FastAPI 封装 /run、/stream 接口，通过 SSE 推送节点级执行事件，Vue3 + "
        "TypeScript 实时渲染执行进度与 Markdown 报告；Docker Compose 一键编排 PostgreSQL、Redis、etcd、MinIO、Milvus "
        "及前后端服务。",
    ]
    for duty in duties:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(duty)
        set_run_font(run, 11, False, None)
        set_numbering(p, 1)

    # 项目成果 ---------------------------------------------------------------
    doc.add_paragraph("项目成果", style="Heading 2")
    metrics = [
        ("回答幻觉率", "由约 25% 降至 6%（200 条内部测评集，人工盲测 + 脚本校验）"),
        ("引用准确率", "达 94%"),
        ("研究完备性满意度", "由 62% 提升至 89%"),
        ("检索阶段耗时", "降低约 35%，双源检索平均响应 < 8s"),
        ("意图路由准确率", "达 96%，减少约 60% 无效深度研究调用"),
        ("低质量信源占比", "由 45% 降至 12%"),
    ]
    table = doc.add_table(rows=1 + len(metrics), cols=2)
    table.autofit = False
    set_table_borders(table)
    set_table_geometry(table, [2700, 6660])

    header_cells = table.rows[0].cells
    for idx, text in enumerate(("指标", "结果")):
        cell = header_cells[idx]
        shade_cell(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        set_run_font(run, 11, True, DARK_BLUE)

    for row_idx, (name, value) in enumerate(metrics, start=1):
        cells = table.rows[row_idx].cells
        for col_idx, text in enumerate((name, value)):
            cell = cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(text)
            set_run_font(run, 11, col_idx == 0, None)

    doc.core_properties.title = "DeepResearch 企业级多 Agent 深度研究智能助手"
    doc.core_properties.author = ""
    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    main()
